import os
from dotenv import load_dotenv, find_dotenv
from docx import Document
from typing import List, Dict
import fitz
import pytesseract
import io
from functools import wraps
from PIL import Image
from newspaper import Article
import string
import re

pytesseract.pytesseract.tesseract_cmd = r"/usr/bin/tesseract"

class DocumentProcessing(object):
    def __init__(self) ->None:
        load_dotenv(find_dotenv())
    
    @staticmethod
    def ExceptionHandeler(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            try:
                return func(*args, **kwargs)
            except (FileNotFoundError, ValueError,Exception,FileNotFoundError, ValueError, NameError, TypeError, KeyError, IndexError, AttributeError, ValueError, AssertionError) as e:
                print(f"An error occurred: {e}")
                raise e
        return wrapper
    
    def datacleaning(self, textfile: str) -> str:
        if not textfile:
            return ""
        
        text = textfile
        text = re.sub(r'\s+', ' ', text.strip())
        text = re.sub(r'[^\x00-\x7F]', ' ', text)
        text = re.sub(r'[-–]', ' ', text) 
        text = re.sub(r'(\w)[|](\w)', r'\1, \2', text)
        
        text = text.replace(" - ", "\n- ").replace(":", ":\n")  

        return text.strip()
    
    @ExceptionHandeler
    def FileProcessing(self, filepath:str) -> str:
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"The file {filepath} does not exist.")
        
        ExtractedData:str = ""
        if filepath.lower().endswith('.docx'):
            doc: Document = Document(filepath)
            for paragraphe in doc.paragraphs:
                ExtractedData += paragraphe.text + '\n'
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        ExtractedData += cell.text + '\n'

            for image in doc.inline_shapes:
                imageData = image._inline.graphic.graphicData.pic.nvPicBlip.blob
                imageData = imageData.decode('utf-8')
                image = Image.open(io.BytesIO(imageData))
                text: str = pytesseract.image_to_string(image)
                ExtractedData += text + '\n'
        
        elif filepath.lower().endswith(('jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff', 'webp')):
            image = Image.open(filepath)
            if filepath.lower().endswith("webp"):
                image = image.convert("RGB")
            text: str = pytesseract.image_to_string(image)
            if text:
                ExtractedData += text
            else:
                raise ValueError("No text found in the image.")
    
        elif filepath.lower().endswith('.txt'):
            with open(filepath, 'r') as file:
                ExtractedData = file.read()
        
        elif filepath.lower().endswith('.pdf'):
            doc: fitz.Document = fitz.open(filepath)
            for page in doc:
                ExtractedData += page.get_text()
            doc.close()
        
        elif filepath.lower().startswith("https://") or filepath.lower().startswith("http://"):
            article = Article(filepath)
            article.download()
            article.parse()
            ExtractedData = article.text
        else:
            raise ValueError("Unsupported file format.")
        
        return self.datacleaning(ExtractedData)